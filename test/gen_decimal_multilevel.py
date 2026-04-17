"""生成一份"十进制多级标题"文档（1 / 1.1 / 1.1.1），用于测试用户手工指定 marker 场景。
文本前缀都是写死的（非自动编号），测分类器在用户提供 H1=1./H2=1.1/H3=1.1.1 时能按意图正确分层。
"""
import random
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

random.seed(31)
OUT = "/home/asus11700f/下载/doc-skill/test/raw_decimal_multilevel.docx"

doc = Document()
p = doc.add_paragraph("关于多级十进制编号的测试文档")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

LONG = "这是一段示例正文。" * 20

for ch in range(1, 4):
    doc.add_paragraph(f"{ch}. 第{ch}章标题")
    for sec in range(1, 4):
        doc.add_paragraph(f"{ch}.{sec} 第{ch}章第{sec}节标题")
        for sub in range(1, 3):
            doc.add_paragraph(f"{ch}.{sec}.{sub} 第{ch}章第{sec}节第{sub}小节标题")
            doc.add_paragraph(LONG)
            doc.add_paragraph(LONG)

doc.save(OUT)
print(f"生成完成：{OUT}  段落 {len(doc.paragraphs)}")
