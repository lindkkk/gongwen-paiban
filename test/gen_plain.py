"""生成一份完全无结构的长文档（没有标题编号、没有目录），测试 fallback 行为。
目标：至少长到几页，正文应被识别为 Body，至少应有 Title（首段）。
"""
from docx import Document
from docx.shared import Pt

OUT = "/home/asus11700f/下载/doc-skill/test/raw_plain.docx"
doc = Document()

# 首段作为主标题候选（短）
p = doc.add_paragraph("关于平台经济治理的一些初步思考")

# 其他都是长段
LONG = "平台经济作为数字时代的重要经济形态，已经深度融入国民经济的各个环节。从就业结构到消费场景，从生产组织到流通方式，平台经济的影响广泛而深远。近年来，围绕平台经济的治理问题引发了学界与实务界的广泛讨论。本文尝试从若干基础维度对这一议题进行初步梳理，以期为后续研究提供线索。"
for _ in range(100):
    doc.add_paragraph(LONG)

doc.save(OUT)
print(f"生成完成：{OUT}")
