"""生成一份：H1/H2/H3 都使用 Word 自动编号 (numPr) 而非手写前缀的测试文档。
H1: numId=1 ilvl=0, chineseCounting, lvlText="%1、"   → 一、二、三、
H2: numId=1 ilvl=1, decimal, lvlText="%2."            → 1. 2. 3.
H3: numId=1 ilvl=2, decimal, lvlText="（%3）"         → （1）（2）（3）
正文里还混了一段用户手工写的"四、"前缀（混合用法），验证两条路径都工作。
"""
import zipfile, os, re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

OUT = "/home/asus11700f/下载/doc-skill/test/raw_auto_multilevel.docx"
TMP = "/tmp/_auto_tmp.docx"

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
def W(t): return f'{{{W_NS}}}{t}'

# 第一步：用 python-docx 生成基础骨架
doc = Document()
doc.add_paragraph("关于多级自动编号的测试文档").alignment = WD_ALIGN_PARAGRAPH.CENTER

# 下面按顺序堆段落，记录每段是什么角色用于后面注入
SPEC = []  # list of (text, role)   role ∈ {H1,H2,H3,Body,H1manual}
SPEC += [('首章起步',                 'H1')]
SPEC += [('首章的第一点说明',         'H2')]
SPEC += [('细分内容示例',             'H3')]
SPEC += [('这是一段普通正文。'*4,    'Body')]
SPEC += [('细分内容示例',             'H3')]
SPEC += [('这是一段普通正文。'*4,    'Body')]
SPEC += [('首章的第二点说明',         'H2')]
SPEC += [('这是一段普通正文。'*4,    'Body')]
SPEC += [('次章推进',                 'H1')]
SPEC += [('次章的第一点说明',         'H2')]
SPEC += [('这是一段普通正文。'*4,    'Body')]
SPEC += [('末章收束',                 'H1')]
# 混入一条手工前缀 H1（测试手工 + 自动混用）
SPEC += [('四、手工前缀的一级标题',  'H1manual')]
SPEC += [('这是一段普通正文。'*4,    'Body')]
# 再来一条中英文括号混搭的 H3（不走 numPr，靠 regex）
SPEC += [('(1)英文括号的三级标题',   'H3manual_en')]
SPEC += [('这是一段普通正文。'*4,    'Body')]
SPEC += [('（2）中文括号的三级标题', 'H3manual_cn')]
SPEC += [('这是一段普通正文。'*4,    'Body')]

for text, role in SPEC:
    doc.add_paragraph(text)

doc.save(TMP)

# 第二步：打开 zip，注入/替换 numbering.xml + 修改 document.xml 给 H1/H2/H3 加 numPr
with zipfile.ZipFile(TMP, 'r') as zin:
    members = {n: zin.read(n) for n in zin.namelist()}

# -- 构造 numbering.xml --
NUMBERING_XML = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W_NS}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="chineseCounting"/>
      <w:lvlText w:val="%1、"/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
    <w:lvl w:ilvl="1">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%2."/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
    <w:lvl w:ilvl="2">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="（%3）"/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1"><w:abstractNumId w:val="0"/></w:num>
</w:numbering>
'''
members['word/numbering.xml'] = NUMBERING_XML.encode('utf-8')

# 确保 document.xml.rels 里有 numbering 的 relationship
rels_xml = members['word/_rels/document.xml.rels'].decode('utf-8')
if 'numbering.xml' not in rels_xml:
    # 在 </Relationships> 前插
    insert = '<Relationship Id="rIdNumbering" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering" Target="numbering.xml"/>'
    rels_xml = rels_xml.replace('</Relationships>', insert + '</Relationships>')
    members['word/_rels/document.xml.rels'] = rels_xml.encode('utf-8')

# 解析 document.xml 给相应段落挂 numPr + 加粗
tree = etree.fromstring(members['word/document.xml'])
body = tree.find(W('body'))
paras = body.findall(W('p'))

# 跳过第一个（标题段），之后按 SPEC 顺序对应
spec_idx = 0
role_to_ilvl = {'H1': 0, 'H2': 1, 'H3': 2}
for p in paras[1:1+len(SPEC)]:
    text, role = SPEC[spec_idx]
    spec_idx += 1

    ppr = p.find(W('pPr'))
    if ppr is None:
        ppr = etree.Element(W('pPr'))
        p.insert(0, ppr)

    if role in role_to_ilvl:
        # 插入 numPr
        numpr = etree.SubElement(ppr, W('numPr'))
        ilvl = etree.SubElement(numpr, W('ilvl'))
        ilvl.set(W('val'), str(role_to_ilvl[role]))
        numid = etree.SubElement(numpr, W('numId'))
        numid.set(W('val'), '1')
        # 加粗（H1/H2/H3 都加）
        for run in p.findall(W('r')):
            rpr = run.find(W('rPr'))
            if rpr is None:
                rpr = etree.Element(W('rPr'))
                run.insert(0, rpr)
            if rpr.find(W('b')) is None:
                etree.SubElement(rpr, W('b'))

    # 手工 H1（"四、..."）同样加粗，但不加 numPr，靠 regex 认
    if role == 'H1manual':
        for run in p.findall(W('r')):
            rpr = run.find(W('rPr'))
            if rpr is None:
                rpr = etree.Element(W('rPr'))
                run.insert(0, rpr)
            if rpr.find(W('b')) is None:
                etree.SubElement(rpr, W('b'))

members['word/document.xml'] = etree.tostring(tree, xml_declaration=True, encoding='UTF-8', standalone=True)

# 写出
if os.path.exists(OUT): os.remove(OUT)
with zipfile.ZipFile(OUT, 'w', zipfile.ZIP_DEFLATED) as zout:
    for name, data in members.items():
        zout.writestr(name, data)
print(f"生成完成：{OUT}")
