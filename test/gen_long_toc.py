"""生成一份 TOC ≥2 页的测试文档，验证 Roman 页码路径。"""
import random
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

random.seed(123)

OUT = "/home/asus11700f/下载/doc-skill/test/raw_long_toc.docx"
doc = Document()

# 封面
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("某某重大专项课题综合研究报告（2026）").font.size = Pt(18)
for _ in range(3): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run("某某研究院").font.size = Pt(14)
doc.add_page_break()

# 超长目录（80 项 → 估算 ≥2 页）
doc.add_paragraph("目录").alignment = WD_ALIGN_PARAGRAPH.CENTER
cn_nums = ["一", "二", "三", "四", "五", "六", "七", "八"]
cn_inner = ["一", "二", "三", "四", "五", "六", "七", "八", "九"]
for ci in range(8):
    doc.add_paragraph(f"{cn_nums[ci]}、第{ci+1}章标题 ..................... {ci*5+1}")
    for si in range(4):
        doc.add_paragraph(f"（{cn_inner[si]}）第{si+1}节内容 ................. {ci*5+si+2}")
        for ki in range(2):
            doc.add_paragraph(f"{ki+1}. 小节标题示例 ................. {ci*5+si+ki+3}")

doc.add_page_break()

# 正文
LONG = "本章围绕核心议题展开系统论述，结合国内外文献与实地调研数据，从理论、方法、实证、政策四个维度进行深入剖析。研究发现，当前相关领域呈现出若干值得关注的新趋势，既有历史延续性又有新的时代特征。"
for ci in range(8):
    doc.add_paragraph(f"{cn_nums[ci]}、第{ci+1}章标题")
    for si in range(4):
        doc.add_paragraph(f"（{cn_inner[si]}）第{si+1}节内容")
        for ki in range(2):
            doc.add_paragraph(f"{ki+1}. 小节标题示例")
            for _ in range(3):
                doc.add_paragraph(LONG)

doc.save(OUT)
print(f"生成完成：{OUT}，段落数 {len(doc.paragraphs)}")
