"""生成一份已经用 Word 内置 Heading 1/2/3 样式标好层级的文档。
测试分类器是否能正确"沿用已有层级"而不是靠文本前缀。
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "/home/asus11700f/下载/doc-skill/test/raw_with_styles.docx"
doc = Document()

# 主标题
p = doc.add_paragraph("数字时代公文处理流程再造研究", style='Title')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()
doc.add_paragraph("目录").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Chapter One Introduction ..... 1")
doc.add_paragraph("Chapter Two Methodology ..... 5")
doc.add_paragraph("Chapter Three Findings ..... 10")

doc.add_page_break()

# 正文：用 Heading 1/2/3 样式，但文本不带中文编号
LONG = "数字化转型已经成为各级政府部门的重要任务，在此过程中，公文作为组织内部与外部沟通的核心载体，其形式与内涵都面临着深刻调整。如何在保持公文严肃性、规范性的同时引入更加高效的信息流转方式，是摆在我们面前的现实课题。"

for ch in ["Introduction", "Methodology", "Findings"]:
    doc.add_paragraph(ch, style='Heading 1')
    for sec in ["Background", "Related Work", "Problem Statement"]:
        doc.add_paragraph(sec, style='Heading 2')
        for sub in ["Context", "Challenges"]:
            doc.add_paragraph(sub, style='Heading 3')
            for _ in range(3):
                doc.add_paragraph(LONG)

doc.save(OUT)
print(f"生成完成：{OUT}")
