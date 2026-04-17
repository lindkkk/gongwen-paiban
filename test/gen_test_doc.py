"""生成一份 ≥50 页的测试 docx。

故意让原稿"杂乱"：
- 有封面（含标题 + 单位 + 日期）
- 有目录（标"目录"二字 + 若干项）
- 一级标题用 "一、二、三、..."  （多于一个，以测试 regex 是否只匹配第一个）
- 二级标题用 "（一）（二）..."
- 三级标题用 "1. 2. 3. ..."
- 混入一个加粗短句（信号兜底测试）
- 原稿字体故意乱设（宋体、黑体随机），正文字号混用小四/四号
- 不设任何 Heading 样式（测试"原稿没标层级"的情况）
"""
import random
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

random.seed(42)

OUT = "/home/asus11700f/下载/doc-skill/test/raw_test.docx"

doc = Document()

# ==== 默认字号 ====
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(11)

# ==== 封面 ====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("关于加快推进产业数字化转型若干问题的研究报告")
run.font.name = '黑体'
run.font.size = Pt(18)
run.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("——技术、制度与生态协同视角").font.size = Pt(14)

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("课题组").font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("2026年4月").font.size = Pt(14)

doc.add_page_break()

# ==== 目录 ====
doc.add_paragraph("目录").alignment = WD_ALIGN_PARAGRAPH.CENTER
toc_items = [
    "一、引言 ..................... 1",
    "（一）研究背景 ............... 2",
    "（二）研究意义 ............... 3",
    "二、文献综述 ................. 5",
    "（一）国内研究现状 ........... 6",
    "（二）国外研究现状 ........... 8",
    "（三）研究述评 ............... 10",
    "三、研究方法与数据 ........... 12",
    "（一）研究方法 ............... 12",
    "（二）数据来源 ............... 14",
    "四、实证分析 ................. 16",
    "（一）描述性统计 ............. 16",
    "（二）回归分析 ............... 20",
    "（三）稳健性检验 ............. 25",
    "五、结论与建议 ............... 28",
    "（一）主要结论 ............... 28",
    "（二）政策建议 ............... 30",
    "六、研究展望 ................. 33",
    "参考文献 ..................... 35",
    "附录 ......................... 38",
]
for item in toc_items:
    doc.add_paragraph(item)

doc.add_page_break()

# ==== 正文 ====
# 用中文数字一级（一、）+ 括号二级（（一））+ 数字三级（1.）方案

def add_raw(text, font_name='宋体', size_pt=11, bold=False):
    """添加一段"原稿"格式的文字。"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    return p

# 一些固定的长段正文，用于堆页数（通用技术/政策类样本文字，不涉及任何特定机构或地域）
LONG_BODY_PARAS = [
    "本报告立足于数字经济快速演进的宏观环境，系统梳理了国内外关于产业数字化转型、信息技术应用、创新生态建设等领域的最新研究成果，并结合实地调研、深度访谈、问卷调查等多种方法，形成了对当前产业数字化转型整体格局的认识。在研究过程中，我们特别注意到，随着云计算、人工智能、大数据等技术的普及，传统产业的生产流程、组织形态、商业模式、价值创造方式均发生着深刻变革，需要从更广阔的视角重新审视产业转型的路径、节奏与支撑体系。",
    "从研究方法上看，本报告采用了定量分析与定性分析相结合的研究策略。在定量层面，我们构建了涵盖技术投入、产出效率、人才结构、资金支持等多维度的评价指标体系，对样本企业进行了系统打分与排序；在定性层面，我们对若干具有代表性的企业开展了半结构化访谈，就其运行机制、组织架构、激励方式等内容进行了深入了解。通过多种方法的交叉验证，我们力图得到更加稳健、可靠的研究结论，避免单一方法可能带来的偏差。",
    "具体而言，我们发现当前产业数字化转型存在若干共性问题：一是高层次复合型人才短缺，尤其是既懂行业又熟悉数字技术的交叉人才严重不足；二是技术应用同质化现象较为突出，部分行业在热点领域重复投入而对真正关键的基础性、前瞻性技术关注不足；三是成果转化机制有待完善，实验室成果与产业化落地之间仍存在明显的对接鸿沟；四是国际交流与国际标准话语权相对薄弱，仍有较大提升空间。",
    "针对上述问题，我们在大量调研的基础上提出了若干对策建议。首先，应出台更有力的复合型人才引进、培养与激励政策，打造具有技术深度与行业广度的骨干队伍；其次，要优化选题机制，鼓励面向真问题、敢于探索长周期、高难度基础研究的攻关项目，并建立相应的容错机制；再次，应搭建更加畅通的产学研协同通道，推动关键成果进入生产环节；最后，要加强国际合作，积极参与全球性议题研究与规则制定，扩大技术与标准的影响力。",
]

# 章节结构，总共6章，每章3-4节，每节2-3小节，每小节3-4段正文 → 粗估约 6*3*3*3=162 段正文，加上封面+目录应当远超50页
chapter_titles = ["引言", "文献综述", "研究方法与数据", "实证分析", "结论与建议", "研究展望"]
section_titles_pool = [
    "研究背景与意义", "国内外研究现状", "核心概念界定", "理论基础",
    "研究方法", "数据来源与处理", "模型构建", "描述性统计",
    "基准回归结果", "异质性分析", "稳健性检验", "作用机制检验",
    "主要结论", "政策建议", "研究不足", "未来展望"
]
subsection_titles_pool = [
    "指标体系构建", "样本选取", "变量定义", "描述统计",
    "相关性分析", "主效应检验", "中介效应", "调节效应",
    "子样本回归", "替换变量", "内生性处理", "稳健性讨论"
]

cn_nums = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
cn_inner = ["一", "二", "三", "四", "五", "六", "七", "八", "九"]

# 目标 ≥50 页。按 A4 / 三号 / 1.5 倍行距估算每页约 30 行 → 需要 ~1500 行内容。
# 设 3-4 段正文 × 5 小节 × 3 节 × 6 章 = 270~360 段正文，每段 ~5 行 = 1350~1800 行。
# 为保险起见每小节 4-5 段。

for i, ch in enumerate(chapter_titles):
    add_raw(f"{cn_nums[i]}、{ch}", font_name='宋体', size_pt=14, bold=True)
    for j in range(3):
        sec_title = random.choice(section_titles_pool)
        add_raw(f"（{cn_inner[j]}）{sec_title}", font_name='宋体', size_pt=12, bold=False)
        for k in range(3):   # 每节 3 小节
            sub_title = random.choice(subsection_titles_pool)
            add_raw(f"{k+1}. {sub_title}", font_name='宋体', size_pt=11, bold=True)
            for _ in range(random.randint(4, 5)):
                add_raw(random.choice(LONG_BODY_PARAS),
                        font_name=random.choice(['宋体', '微软雅黑', '楷体']),
                        size_pt=random.choice([10, 11, 12]),
                        bold=False)

# 混入一段看起来像标题的加粗短句（测试信号兜底）
add_raw("附加观察", font_name='黑体', size_pt=12, bold=True)
add_raw("在研究过程中我们还注意到若干值得后续深入挖掘的观察，这些观察虽然未在正文中展开但可能对后续研究具有一定的启发意义。", font_name='宋体', size_pt=11)

# 结尾
add_raw("参考文献", font_name='宋体', size_pt=14, bold=True)
for i in range(20):
    add_raw(f"[{i+1}] 某某作者. 某某文献标题[J]. 某某期刊, 2025, (1): 1-10.", font_name='宋体', size_pt=10)

doc.save(OUT)

# 统计
from docx import Document as D
d = D(OUT)
print(f"生成完成：{OUT}")
print(f"段落数：{len(d.paragraphs)}")
