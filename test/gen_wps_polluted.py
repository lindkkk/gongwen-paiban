"""生成一份"WPS 风格"污染过的 docx，复现用户那份真实文件的 pollution 模式：
- rFonts 带 *Theme 主题字体属性（导致 Word 用主题覆盖显式字体 → 最终显示宋体）
- 二级标题挂 numPr 自动编号列表（导致缩进错乱）
- pPr 里塞 kinsoku/snapToGrid/边框/段落标记 rPr/ 等各种残留
用来验证 v2 格式化能彻底清掉这些污染。
"""
import random, zipfile, shutil, os, re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

random.seed(77)
OUT = "/home/asus11700f/下载/doc-skill/test/raw_wps_polluted.docx"
TMP = "/tmp/_wps_tmp.docx"

# 先用 python-docx 生成干净的结构
doc = Document()
# 主标题
p = doc.add_paragraph("贸易正在发生变化")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.size = Pt(18)

doc.add_paragraph("——2026年达沃斯论坛的10个深刻见解")

HEADINGS = [
    "紧张局势起伏不定，格陵兰关税取消",
    "制造业在复苏",
    "数字贸易的新边疆",
    "供应链弹性成为竞争优势",
    "服务贸易占比持续上升",
    "绿色关税成为谈判焦点",
    "新兴市场交易加速",
    "知识产权保护趋严",
    "金融服务的数字化转型",
    "全球化进入新范式",
]

LONG = "全球贸易格局正在经历自二战以来最深刻的一次重塑，地缘政治、技术变革、气候议程三股力量交织作用，使得传统的贸易理论和国别竞争力评估框架都面临重大挑战。一方面，主要经济体的保护主义倾向明显抬头，关税与非关税壁垒层层加码；另一方面，数字技术使得跨境服务贸易以前所未有的速度扩张，对既有的规则体系形成持续冲击。"

for i, h in enumerate(HEADINGS):
    doc.add_paragraph(h)
    for _ in range(3):
        doc.add_paragraph(LONG)

doc.save(TMP)

# 读出来，注入污染
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NSMAP = {'w': W_NS}
def W(t): return f'{{{W_NS}}}{t}'

# 污染用的 pPr/rPr 模板
POLLUTED_PPR_JUNK = """
<w:keepNext xmlns:w="{ns}" w:val="0"/>
<w:keepLines xmlns:w="{ns}" w:val="0"/>
<w:pageBreakBefore xmlns:w="{ns}" w:val="0"/>
<w:widowControl xmlns:w="{ns}"/>
<w:suppressLineNumbers xmlns:w="{ns}" w:val="0"/>
<w:pBdr xmlns:w="{ns}">
  <w:top w:val="none" w:color="auto" w:sz="0" w:space="0"/>
  <w:left w:val="none" w:color="auto" w:sz="0" w:space="0"/>
  <w:bottom w:val="none" w:color="auto" w:sz="0" w:space="0"/>
  <w:right w:val="none" w:color="auto" w:sz="0" w:space="0"/>
</w:pBdr>
<w:kinsoku xmlns:w="{ns}"/>
<w:wordWrap xmlns:w="{ns}"/>
<w:overflowPunct xmlns:w="{ns}"/>
<w:topLinePunct xmlns:w="{ns}" w:val="0"/>
<w:autoSpaceDE xmlns:w="{ns}"/>
<w:autoSpaceDN xmlns:w="{ns}"/>
<w:bidi xmlns:w="{ns}" w:val="0"/>
<w:adjustRightInd xmlns:w="{ns}"/>
<w:snapToGrid xmlns:w="{ns}"/>
<w:textAlignment xmlns:w="{ns}" w:val="auto"/>
<w:rPr xmlns:w="{ns}">
  <w:rFonts w:hint="eastAsia" w:asciiTheme="minorEastAsia" w:hAnsiTheme="minorEastAsia" w:eastAsiaTheme="minorEastAsia" w:cstheme="minorEastAsia"/>
  <w:color w:val="000000"/>
  <w:sz w:val="21"/>
  <w:szCs w:val="21"/>
</w:rPr>
""".format(ns=W_NS).strip()

POLLUTED_RPR_THEME = f"""
<w:rFonts xmlns:w="{W_NS}" w:hint="eastAsia" w:asciiTheme="minorEastAsia" w:hAnsiTheme="minorEastAsia" w:eastAsiaTheme="minorEastAsia" w:cstheme="minorEastAsia"/>
""".strip()

# 把 TMP 打开改 document.xml
with zipfile.ZipFile(TMP, 'r') as zin:
    docxml = zin.read('word/document.xml')
    names = zin.namelist()
    members = {n: zin.read(n) for n in names}

tree = etree.fromstring(docxml)
body = tree.find(W('body'))
paragraphs = body.findall(W('p'))

heading_text_set = set(HEADINGS)

for pi, p in enumerate(paragraphs):
    text = ''.join([t.text or '' for t in p.iter(W('t'))])

    # 确保 pPr 存在
    ppr = p.find(W('pPr'))
    if ppr is None:
        ppr = etree.SubElement(p, W('pPr'))
        p.insert(0, ppr)

    # 注入 pPr 污染（放在既有子元素前）
    junk_fragment = etree.fromstring('<w:root xmlns:w="{}">{}</w:root>'.format(W_NS, POLLUTED_PPR_JUNK))
    for i, junk in enumerate(list(junk_fragment)):
        ppr.insert(i, junk)

    # 如果是二级标题文本，加 numPr + OutlineLevel（模拟 WPS 自动编号 + 大纲层级）
    if text in heading_text_set:
        numpr = etree.SubElement(ppr, W('numPr'))
        ilvl = etree.SubElement(numpr, W('ilvl'))
        ilvl.set(W('val'), '0')
        numid = etree.SubElement(numpr, W('numId'))
        numid.set(W('val'), '1')
        # OutlineLevel=1 表示二级标题
        outl = etree.SubElement(ppr, W('outlineLvl'))
        outl.set(W('val'), '1')
        # 给 run 加 bold（模拟真实二级标题的加粗）
        for run in p.findall(W('r')):
            rpr = run.find(W('rPr'))
            if rpr is None:
                rpr = etree.SubElement(run, W('rPr'))
                run.insert(0, rpr)
            if rpr.find(W('b')) is None:
                bold = etree.SubElement(rpr, W('b'))

    # run 的 rPr 里注入主题字体
    for run in p.findall(W('r')):
        rpr = run.find(W('rPr'))
        if rpr is None:
            rpr = etree.SubElement(run, W('rPr'))
            run.insert(0, rpr)
        # 插入主题字体 rFonts（取代或新增）
        existing_fonts = rpr.find(W('rFonts'))
        if existing_fonts is not None:
            existing_fonts.set(W('asciiTheme'), 'minorEastAsia')
            existing_fonts.set(W('hAnsiTheme'), 'minorEastAsia')
            existing_fonts.set(W('eastAsiaTheme'), 'minorEastAsia')
            existing_fonts.set(W('cstheme'), 'minorEastAsia')
            existing_fonts.set(W('hint'), 'eastAsia')
        else:
            theme_fonts = etree.fromstring(POLLUTED_RPR_THEME)
            rpr.insert(0, theme_fonts)

# 写出
docxml_new = etree.tostring(tree, xml_declaration=True, encoding='UTF-8', standalone=True)
members['word/document.xml'] = docxml_new

if os.path.exists(OUT):
    os.remove(OUT)
with zipfile.ZipFile(OUT, 'w', zipfile.ZIP_DEFLATED) as zout:
    for name, data in members.items():
        zout.writestr(name, data)

print(f"生成完成：{OUT}")
# 打印前几个段落的 pPr 确认污染已注入
tree2 = etree.fromstring(docxml_new)
body2 = tree2.find(W('body'))
for p in body2.findall(W('p'))[:3]:
    text = ''.join([t.text or '' for t in p.iter(W('t'))])[:30]
    has_numpr = p.find('.//' + W('numPr')) is not None
    fonts = p.find('.//' + W('rPr') + '/' + W('rFonts'))
    has_theme = False
    if fonts is not None:
        has_theme = any(fonts.get(W(a)) for a in ('asciiTheme','hAnsiTheme','eastAsiaTheme','cstheme'))
    print(f"  text={text!r} numPr={has_numpr} themeFont={has_theme}")
